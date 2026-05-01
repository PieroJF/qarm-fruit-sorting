"""
Build final Applied Robotics report following coursework spec.
Output: applied_robotics_report_FINAL.docx in Test_result/.

Structure (per Coursework PDF):
  Title page (1)
  PART 1 [12 pages]: Ch 1-5 (Intro, Background, Methodology, Simulation, Control)
  Soft Gripper [2 pages]: Ch 6
  PART 2 [7 pages]: Ch 7-9 (Validation, Discussion, Conclusion)
  Feedback Reflection [1 page]: Ch 10

English level: B2 for new content; preserved existing content as written.
"""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, 'applied_robotics_report_FINAL.docx')

doc = Document()

# ---------- page setup ----------
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
pf = style.paragraph_format
pf.line_spacing = 1.15
pf.space_after = Pt(6)

# headings
for hname, sz in [('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11)]:
    h = doc.styles[hname]
    h.font.name = 'Times New Roman'
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)


def H1(t):
    p = doc.add_heading(t, level=1)
    return p


def H2(t):
    p = doc.add_heading(t, level=2)
    return p


def H3(t):
    p = doc.add_heading(t, level=3)
    return p


def P(text, italic=False, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def BLANK():
    doc.add_paragraph('')


def PAGEBREAK():
    doc.add_page_break()


def FIG(filename, caption, width_in=5.5):
    path = os.path.join(HERE, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Inches(width_in))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(10)


# ====================================================================
# TITLE PAGE
# ====================================================================
for _ in range(4):
    BLANK()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('UNIVERSITY OF BIRMINGHAM')
r.bold = True
r.font.size = Pt(20)

BLANK()
BLANK()
BLANK()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Robotic Pick and Place for Fruit Sorting')
r.bold = True
r.font.size = Pt(26)

BLANK()
BLANK()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Final Group Report')
r.italic = True
r.font.size = Pt(14)

BLANK()
BLANK()
BLANK()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Applied Robotics (04 39984)')
r.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Module Coordinator: Dr Amir Hajiyavand')
r.font.size = Pt(13)

BLANK()
BLANK()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Group Members')
r.bold = True
r.font.size = Pt(12)

for name in ['Piero Flores', 'Yichang Chao', 'Ran Zhang', 'Zihen Huang']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.font.size = Pt(12)

BLANK()
BLANK()
BLANK()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Submission date: 1 May 2026')
r.italic = True
r.font.size = Pt(11)

PAGEBREAK()

# ====================================================================
# PART 1 — DESIGN AND CONTROL STRATEGY (12 pages)
# ====================================================================

# --- 1. Introduction ---
H1('1. Introduction')

H2('1.1 Background and Motivation')
P('Pick-and-place is one of the most common applications of robotic manipulators in '
  'industry. The agricultural sector in particular faces rising labour costs and a need '
  'for consistent product handling, and automated fruit sorting is one realistic place '
  'where robotic manipulation can replace repetitive manual work. Soft handling is '
  'especially important for fresh produce because rigid grippers can bruise the fruit '
  'and lower its market value.')
P('This module asks us to design and validate a complete pick-and-place solution on a '
  'small 4-DOF educational arm, the Quanser QArm. The arm is paired with an Intel '
  'RealSense D415 RGB-D camera, which gives both colour and depth information. Working '
  'with this hardware lets us experience the full pipeline from sensor data to actuator '
  'command, instead of working only at the algorithm level.')

H2('1.2 Project Objectives')
P('The system has to identify three classes of fruit (banana, tomato, strawberry) '
  'placed at random in the workspace and sort them into three labelled baskets. The '
  'fruit set is fixed at three bananas, five tomatoes and a six-pack of strawberries, '
  'as specified by the coursework brief. There are two operational modes that the '
  'system must support:')
P('• Fully automated mode: the robot detects, locates, picks, and places fruit without '
  'human input.')
P('• Remote-controlled mode: a human operator drives the same robot through the same '
  'task using a graphical user interface.')
P('The deliverable is a working physical demonstration on the actual QArm hardware, a '
  'soft-gripper design with an analytical safety analysis, and this technical report '
  'covering design, simulation, and experimental validation.')

H2('1.3 Report Structure')
P('This report follows the structure laid out in the coursework brief. Part 1 (Chapters '
  '1–5) covers the background, the design of the kinematics and control stack, the '
  'Simulink simulation work, and the implementation of both operational modes. Chapter '
  '6 is the soft-gripper design and analysis. Part 2 (Chapters 7–9) covers the '
  'experimental validation on the physical QArm, including the four-mode test '
  'campaign, the failure analysis, the discussion of discrepancies, and the conclusion. '
  'Chapter 10 is a one-page reflection on the feedback we received from the lecturers '
  'and the PGTAs during the lab sessions.')

# --- 2. Background and Literature Review ---
H1('2. Background and Literature Review')

H2('2.1 Robotic Pick-and-Place in Industry')
P('Pick-and-place is the canonical task of industrial manipulators and has been widely '
  'deployed in electronics assembly, food packaging, and warehouse automation. The '
  'basic capability has three layers: a perception layer that locates objects in the '
  'workspace, a planning layer that decides the order and trajectory of moves, and a '
  'control layer that executes the joint commands while respecting the hardware '
  'envelope. Industrial systems usually run these on 6-DOF or 7-DOF arms with vision '
  'systems calibrated to the work cell, but the same three-layer structure also '
  'applies to small educational platforms such as the QArm.')
P('Two control architectures are common in the literature. Position-mode controllers '
  'send joint targets at a fixed cycle rate and rely on the on-board servo loop to '
  'track them. Torque-mode controllers send torque commands directly and run the '
  'tracking loop at the host. Position mode is simpler and is sufficient for '
  'pick-and-place because the dynamic envelope is well below the actuator limits. We '
  'use position mode throughout this project.')

H2('2.2 Vision-Based Fruit Detection')
P('Two main families of approaches exist for detecting fruit in cluttered scenes. The '
  'classical computer vision family uses colour-space segmentation (typically HSV or '
  'Lab), morphological operations, and contour-based shape analysis. This family is '
  'data-light, computationally cheap, and easy to debug, but the thresholds need to be '
  'tuned to the lighting and the camera. The deep learning family uses convolutional '
  'networks (e.g. YOLO, Mask R-CNN) trained on annotated images. This family handles '
  'non-canonical orientations and occlusion better, but requires labelled data, '
  'training time, and GPU resources.')
P('We chose the classical approach because the lab lighting is stable, the fruit set '
  'is fixed and small, and we did not have the time or labelled data needed to train '
  'a network. The classical approach also makes every step inspectable, which was '
  'important during debugging. The detection logic is described in Section 5.1.')

H2('2.3 Soft Gripper Approaches')
P('Three soft-gripper paradigms are common in the literature. Pneumatic networks '
  '(PneuNets) use compressed-air channels embedded in silicone to bend on demand. '
  'Granular jamming uses a flexible bag filled with coffee grounds or similar; the '
  'bag conforms to the object and is then evacuated to lock the shape. Fin Ray '
  'fingers use a triangular truss structure that deflects passively in proportion to '
  'the contact force. Each has different trade-offs in actuation complexity, payload, '
  'and how easy it is to analyse on paper.')
P('We selected the Fin Ray approach for this project. The reasoning, the analytical '
  'model, and the safety-factor analysis are presented in Chapter 6.')

# --- 3. System Design and Methodology ---
H1('3. System Design and Methodology')

H2('3.1 System Architecture and Toolchain Decision')
P('Our original control plan for the QArm was based on the standard Quanser workflow, '
  'in which the controller is built as a Simulink model and deployed to the robot '
  'through the QUARC real-time toolchain. During the first lab session after the '
  'Easter break, we discovered that none of the workstations available to our group '
  'had an active QUARC license. The Simulink build step failed at the deployment '
  'stage with a license error, and the same error reappeared on every machine we '
  'tried. Because the rest of our planned pipeline, including the trajectory '
  'generator, the gripper command path, and the vision-to-pick handover, all assumed '
  'that QUARC would manage the real-time loop, the missing license blocked the '
  'project at its lowest layer. Without a working way to send any joint command to '
  'the hardware, we could not even verify that the arm was wired correctly, let '
  'alone start testing the higher-level sorting logic.')
P('After confirming that the failure was a licensing issue rather than a hardware or '
  'installation fault, we raised the problem with one of the lab teaching assistants '
  'during the next session. The TA explained that QUARC is one of two supported ways '
  'to drive the QArm and is not strictly required for our use case. The other path '
  'is the Quanser Python SDK, a library that exposes the same low-level hardware '
  'functions QUARC uses internally but makes them available through a plain Python '
  'API. The SDK ships with the lab installation by default, runs on a standard '
  'Python interpreter, and does not need a QUARC seat in order to read sensor data '
  'or send actuator commands.')
P('Acting on this advice, we redesigned the control stack as a pure Python pipeline '
  'built on top of the SDK. A thin driver layer reads joint angles, end-effector '
  'position, and gripper state directly from the QArm card, and writes commanded '
  'joint targets and gripper setpoints back to the same card at a fixed cycle rate. '
  'All higher-level logic, including forward and inverse kinematics, trajectory '
  'interpolation, vision processing, and the sorting state machine, was moved into '
  'Python rather than Simulink. Simulink was kept as a presentation layer for the '
  'lab demonstrations and was given no responsibility for real-time control. With '
  'this change in place, the entire system became license-independent. Every result '
  'reported in this document was produced under the SDK pipeline, on a standard lab '
  'machine, without any QUARC activation.')
P('The architecture is therefore a "Simulink-as-facade" design. The MATLAB Function '
  'blocks in the Simulink slices call into Python via py.* using coder.extrinsic. '
  'This keeps the lab deliverables (the .slx models that the assessment expects) '
  'while moving all execution into a path that we can actually run.')

H2('3.2 Fundamentals of Robotics: Forward and Inverse Kinematics')
P('This section establishes the mathematical foundation for the fruit-sorting '
  'pipeline. Higher-level behaviours, including vision-driven picking, state-machine '
  'sequencing, and teleoperation, all reduce to finding the end-effector position '
  'for a given joint configuration (forward kinematics) and finding the joint '
  'configuration for a commanded Cartesian pose (inverse kinematics). The Quanser '
  'QArm used in this project is a 4-degree-of-freedom serial manipulator with a '
  'two-finger gripper. Its limited redundancy makes an analytical treatment of these '
  'problems feasible and necessary for real-time control.')

H3('3.2.1 Kinematic Structure and DH Parameters')
P('The QArm is modelled as an open kinematic chain of four revolute joints denoted '
  'φ₁ … φ₄, corresponding physically to the base yaw, shoulder pitch, elbow pitch, '
  'and wrist roll. Because the mechanical elbow has a small structural offset '
  '(L₃ = 0.05 m) that is not collinear with the upper arm (L₂ = 0.35 m), the '
  'manipulator does not map directly onto a textbook planar 3R arm. To avoid '
  'carrying this offset through later derivations, we follow Quanser\'s recommended '
  'lumped-parameter convention and fold the upper-arm link and the offset into a '
  'single effective length:')
P('λ₂ = √(L₂² + L₃²) = 0.3536 m,    β = arctan(L₃ / L₂) = 0.1419 rad (≈ 8.13°)')
P('while the base height and the concatenated forearm-plus-end-effector segment are '
  'kept as λ₁ = L₁ = 0.14 m and λ₃ = L₄ + L₅ = 0.40 m. The geometric bias angle β is '
  'absorbed into the joint-angle offsets that relate the physical joint readings φ '
  'to the DH-frame angles θ:')
P('θ₁ = φ₁,    θ₂ = φ₂ − π/2 + β,    θ₃ = φ₃ − β,    θ₄ = φ₄')
P('The standard Denavit–Hartenberg table is then constructed with a₁ = 0, '
  'α₁ = −π/2, d₁ = λ₁; a₂ = λ₂, α₂ = 0, d₂ = 0; a₃ = 0, α₃ = −π/2, d₃ = 0; and '
  'a₄ = 0, α₄ = 0, d₄ = λ₃. The mechanical joint limits, enforced both in software '
  'clamps (qarm_kinematics.py, JOINT_LIMITS) and in the IK solver\'s validity '
  'check, are ±170° for φ₁ and φ₄ and ±90° for φ₂ and φ₃. These limits shape the '
  'reachable workspace into a truncated hemispherical shell of approximate outer '
  'radius 0.75 m, consistent with the workspace mapping reported in the interim '
  'deliverable (maximum horizontal reach 0.7536 m, maximum height 0.8936 m).')

H3('3.2.2 Forward Kinematics')
P('Given a physical joint vector φ, the end-effector pose in the base frame is the '
  'ordered product of four homogeneous transforms derived from the DH table:')
P('T⁰₄(θ) = A₁(θ₁) · A₂(θ₂) · A₃(θ₃) · A₄(θ₄)')
P('Expanding the product yields the closed-form position equations:')
P('p_x = λ₂ cos θ₁ cos θ₂ − λ₃ cos θ₁ sin(θ₂ + θ₃)')
P('p_y = λ₂ sin θ₁ cos θ₂ − λ₃ sin θ₁ sin(θ₂ + θ₃)')
P('p_z = λ₁ − λ₂ sin θ₂ − λ₃ cos(θ₂ + θ₃)')
P('At the home configuration φ = [0, 0, 0, 0], the forward model returns '
  'p = [0.450, 0.000, 0.490] m, which matches the teach-pendant readout on the lab '
  'QArm to machine precision. Two practical points are worth noting. First, because '
  'the manipulator only has four joints, the reachable orientations form a '
  'one-parameter family for a given target position. The wrist roll θ₄ is largely '
  'independent of position, so the IK routine exposes it as a user-settable '
  'argument γ rather than solving for it. Second, the forward map is also evaluated '
  'inside the inverse solver for residual computation and for finite-difference '
  'Jacobian assembly.')

H3('3.2.3 Inverse Kinematics: Analytical Seed plus Newton–Raphson Refinement')
P('For a 4-DOF arm, the inverse kinematics is geometrically well-posed only for the '
  'position sub-problem (three equations and three free joints once γ is pinned). '
  'We therefore decompose IK into two phases: a closed-form analytical seed, '
  'followed by a damped least-squares correction that absorbs any residual '
  'geometric inconsistency and compensates for numerical drift near singularities.')
P('The analytical seed phase begins by projecting the target p = (p_x, p_y, p_z) '
  'into the vertical plane that contains the base axis, reducing the problem to a '
  '2R planar inverse kinematics in the variables (D, H), where '
  'D = ±√(p_x² + p_y²) is the horizontal reach and H = λ₁ − p_z is the vertical '
  'drop from the shoulder joint. Applying the half-angle substitution to the '
  'standard two-link law of cosines yields four candidate solutions for (θ₂, θ₃), '
  'corresponding to the combinations of left-reach or right-reach with elbow-up or '
  'elbow-down configurations:')
P('θ₃ = 2 arctan2(−C ± √(C² + F²), F),    F = (D² + H² − A² − C²) / (2A)')
P('with A = λ₂ and C = −λ₃. For each (θ₂, θ₃) pair, the base yaw is recovered as '
  'θ₁ = arctan2(p_y/δ, p_x/δ) where δ = λ₂ cos θ₂ − λ₃ sin(θ₂ + θ₃). All four '
  'candidates are converted back to physical angles, wrapped into [−π, π], and '
  'screened against the joint limits. Among the surviving configurations, the one '
  'minimising ‖φ‖² is selected. This cost pushes the solver toward compact postures '
  'that keep the arm away from singularities and reduce unnecessary sweeping '
  'motions between pick and place.')
P('To guarantee Cartesian accuracy at the tolerance of 10⁻⁶ m enforced by the '
  'convergence check, and to absorb residual drift near the workspace boundary '
  'where the discriminant C² + F² approaches zero, a Newton–Raphson update is '
  'applied in task space:')
P('Δφ = Jᵀ (J Jᵀ + λI)⁻¹ (p_target − p_current)')
P('where J ∈ ℝ³ˣ⁴ is a three-point-stencil finite-difference Jacobian computed '
  'with a step size of 10⁻⁶, and λ = 10⁻³ is a Levenberg–Marquardt damping factor '
  'that suppresses the near-singular inversion when the arm is close to full '
  'extension. Each iterate is re-clamped to the joint limits, and the wrist is '
  'held at the user-specified γ so that it does not drift as a free variable. '
  'Offline validation in validate_python.py exercises the solver on six '
  'representative targets that span the three basket positions, the approximate '
  'workspace extremes, and the home pose. In every case the solver converges '
  'inside the 50-iteration budget and satisfies the 10⁻⁶ m Cartesian tolerance.')

H2('3.3 Trajectory Planning')
P('Once IK has produced a feasible joint vector for each waypoint, the motion '
  'between consecutive waypoints must be generated as a time-parameterised '
  'trajectory that the QArm can track at 500 Hz without inducing audible or '
  'mechanical jerk. Two polynomial blending schemes are implemented in '
  'trajectory.py. The autonomous FSM uses the quintic form by default, while the '
  'cubic form is retained as a baseline and is used for joint-space smooth_move '
  'transitions.')
P('The cubic baseline imposes zero velocity at t = 0 and t = T:')
P('p(t) = p_s + (3 Δp / T²) t² − (2 Δp / T³) t³,    Δp = p_e − p_s')
P('Differentiation yields a symmetric bell-shaped velocity profile peaking at '
  '1.5 Δp / T at the mid-segment, and a piecewise-linear acceleration that is '
  'finite but discontinuous at segment boundaries; this discontinuity reaches the '
  'drivetrain as an impulsive jerk every time a new segment begins.')
P('To suppress this boundary jerk, the planner used at runtime '
  '(quintic_trajectory in trajectory.py, called from sorting_controller.py:337) '
  'raises the polynomial order to five, with zero velocity and zero acceleration '
  'imposed at both endpoints:')
P('p(t) = p_s + Δp · s(τ),    s(τ) = 10τ³ − 15τ⁴ + 6τ⁵,    τ = t / T')
P('The resulting acceleration profile is a single symmetric S-curve that vanishes '
  'at both ends, so the jerk stays bounded across segment concatenations. The unit '
  'test test_quintic_boundary_conditions in test_integration.py verifies that '
  'p(0) = p_s, p(T) = p_e, and that both first and second numerical derivatives '
  'are below 10⁻³ at the boundaries. This property is structurally impossible for '
  'the cubic baseline and is practically relevant because the geared drivetrain '
  'of the QArm produces audible noise when jerk is high.')
P('Segment durations are fixed constants inside sorting_controller.py and shared '
  'with the simulation twin in sorting_controller_sim.py: T_TRANSIT = 2.0 s for '
  'long home-to-approach and basket transits, T_APPROACH = 1.0 s for the '
  'intermediate approach and basket-above moves, T_PICK = 0.8 s for the final '
  'vertical descent, and T_DWELL = 0.5 s for the quasi-static hold while the '
  'gripper actuates. The 2-second transit choice was inherited from previous '
  'experimental benchmarks, where shorter values drove the commanded peak '
  'velocity and acceleration into actuator saturation.')

# --- 4. Simulation and Simulink Integration ---
H1('4. Simulation and Simulink Integration')

H2('4.1 Vertical Slicing Approach')
P('Because the real-time loop runs in Python (see Section 3.1), the role of '
  'Simulink in this project is to verify that the data flow between the MATLAB '
  'side and the Python side is correct, and to provide the .slx deliverables '
  'expected by the assignment. We adopted a "vertical slicing" strategy: each '
  'slice is a small Simulink model that exercises one specific data path '
  'end-to-end. Once each slice is verified, they are integrated into the '
  'top-level model FruitSorting_Hardware.slx.')

H2('4.2 MATLAB–Python Bridge for Inverse Kinematics')
P('The model Fruitsorting_slice.slx verifies whether the inverse-kinematics '
  'bridge between Simulink and Python functions correctly. A target Cartesian '
  'coordinate [0.30; −0.20; 0.05] is fed into the input port p of the IK block. '
  'The block calls qarm_kinematics.py, which uses the DH parameters of the '
  'manipulator together with the analytical solver and the Newton–Raphson '
  'refinement (Section 3.2.3) to obtain the four joint angles. A 4×1 array is '
  'produced at the output port phi and routed through output port 1 into a '
  'To Workspace sink, so we can confirm the four joint angles are returned to '
  'the Simulink environment without any data loss.')

H2('4.3 Vision Link Verification')
P('The model Fruitsorting_slice_vision.slx validates the vision module and '
  'tracks the flow of image-recognition data. The Trigger block emits a '
  'one-shot signal that requests a single detection. Once the Detect block '
  'receives this signal, it invokes the background Python script '
  'fruit_detector.py, which reads a test image and processes it through HSV '
  'colour and shape recognition. The Detect block then outputs an 8×5 data '
  'matrix together with a scalar count. The matrix stores the identified fruit '
  'IDs along with their xyz coordinates, while the scalar reports how many '
  'fruits are actually detected. The results are archived in out.det_log and '
  'out.count_log, and both are forwarded to the MATLAB base workspace.')

H2('4.4 Self-Driven Open-Loop FSM Test')
P('The model Fruitsorting_slice_fsm.slx checks whether the FSM logic on the '
  'Python side completes the entire sorting procedure as expected. At every '
  'simulation step, an integrator accumulates a constant input of 1, producing '
  'a monotonically increasing integer sequence. This sequence serves as the '
  'tick signal supplied to the FSM and provides a discrete time base. When '
  'tick equals zero, a reset pulse is sent to the Python side so that the '
  'state machine is reinitialised, guaranteeing a clean initial state at the '
  'start of each test.')
P('The FSM block bridges and invokes the controller script on the Python '
  'side. It takes the current tick and the joint angles from the previous '
  'cycle, phi_prev, as its runtime inputs. After the internal logic '
  'evaluates these inputs, the block emits four control signals: the joint '
  'command for the next cycle, phi_cmd; the gripper control signal, gripper; '
  'the current task-stage code, state_id; and a Boolean flag, done, that '
  'indicates whether the procedure has finished. Four To Workspace blocks '
  'export these commands and state variables in real time and store them in '
  'the MATLAB base workspace. Once the simulation terminates, these '
  'variables are retrieved and plotted as timing diagrams, so we can assess '
  'whether the system states match expectations.')
P('A phi_delay block is also introduced so that the command issued in the '
  'current step, phi_cmd, is delayed by one time step and fed back to the '
  'FSM in the next cycle as the previous position, phi_prev. This emulates '
  'the state update that follows hardware execution and allows the state '
  'machine to advance continuously through the full sorting procedure.')

H2('4.5 Top-Level Hardware-in-the-Loop Integration')
P('The model FruitSorting_Hardware.slx integrates the previously validated '
  'kinematics, vision, and FSM components, and connects them to the low-level '
  'hardware I/O. The FSM block receives the current tick together with the '
  'measured joint positions of the manipulator, joints_fb. Combining these '
  'signals with the data from the vision module, it computes where the arm '
  'should move next, phi_cmd, whether the gripper should open or close, '
  'gripper, and broadcasts the current task progress through state_id and '
  'done. The QarmIO block consumes the motion commands issued by the FSM, '
  'namely phi_cmd and gripper_cmd, and reads the mode switch qmode. After '
  'each command is executed, the block returns the outcome to the rest of '
  'the model, including the actual joint angles read back from the motor '
  'encoders, joints_fb, and the gripper position. Five logging sinks, '
  'including out.phi_cmd_log, record every command and feedback signal '
  'produced during execution and store them in MATLAB memory for later '
  'scoring or trajectory replay.')

# --- 5. Control System Implementation ---
H1('5. Control System Implementation')

H2('5.1 Vision Module')

H3('5.1.1 Overall Architecture')
P('The vision module is the core perception subsystem of the fruit-sorting '
  'robotic arm. Its job is to perform fruit recognition, pose estimation, and '
  'base-frame coordinate projection on the RGB-D images provided by the Intel '
  'RealSense D415 camera. The module is built on classical computer-vision '
  'techniques. By combining HSV colour segmentation, morphological operations, '
  'geometric feature filtering, and solvePnP for extrinsic estimation, we '
  'built a processing pipeline that is light in computation and easy to '
  'interpret, while avoiding the training data and computing resources that '
  'a deep-learning model would require.')
P('The whole vision module is started from main_final.py. Inside, it '
  'connects camera.py, survey_capture.py, fruit_detector.py and '
  'picker_viewer.py in order, while session_cal.py is used to manage the '
  'calibration data in one place. Before the system starts, the offline '
  'calibration toolchain made up of calibrate_chessboard.py, '
  'calibrate_extrinsics.py and homography_solver.py must be run once to '
  'produce session_cal.json, which is then used as the calibration source '
  'at runtime.')
P('On the hardware side, camera.py opens the D415 camera through the '
  'Video3D interface of the Quanser SDK. It starts a 1280×720 BGR colour '
  'stream and a uint16 depth stream in millimetres at the same time. '
  'Auto-exposure, auto white balance and the IR emitter are all turned on '
  'during initialisation, so that later HSV segmentation will not break '
  'when the lighting changes. The camera intrinsics (fx, fy, cx, cy) are '
  'read directly from the sensor when possible, and only fall back to '
  'hard-coded default values when the read fails. This step is important '
  'for the accuracy of the later pixel-to-base-frame projection.')
P('After a BGR frame is captured, the pipeline first calls '
  'cv2.cvtColor(bgr, cv2.COLOR_BGR2HSV) to convert the image into the HSV '
  'colour space. Compared to RGB, HSV puts the hue into a single channel, '
  'which makes it more robust to brightness changes and therefore better '
  'suited as the basis for colour-threshold segmentation. After that, the '
  'system builds an HSV mask for each fruit class using cv2.inRange() to '
  'keep the pixels that fall in the range. For red targets whose hue '
  'crosses the H = 0/180 seam (tomato and strawberry body), two inRange '
  'calls are merged with cv2.bitwise_or. The raw mask still has small '
  'spots and holes from edge highlights and background noise, so a '
  'MORPH_OPEN (5×5 kernel) is applied to remove small noise points and a '
  'MORPH_CLOSE (9×9 kernel) is applied to fill the holes inside the '
  'fruit, keeping the contour boundaries complete.')
P('The cleaned mask is then sent to '
  'cv2.findContours(mask, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE) to get the '
  'outermost contours. Each contour must pass the geometric and colour '
  'gates of its class to become a valid detection candidate. The '
  'candidates that pass the gates are then sampled with '
  'sample_depth_at_pixel over a 5×5 patch to estimate the height of the '
  'fruit\'s top surface using the median depth. When the number of valid '
  'samples is too small, or the values fall outside a reasonable range, '
  'the system falls back to the per-class default heights stored in '
  '_FRUIT_TOP_Z_MM. Finally, pixel_to_base_frame combines the pixel and '
  'depth into a camera ray, rotates it into the base frame using the '
  'extrinsics solved by solvePnP, and intersects it with the target '
  'plane to obtain the (X, Y, Z) coordinate. The detector itself is '
  'stateless: each frame is processed independently, which gives the '
  'system good real-time performance.')

H3('5.1.2 HSV Thresholds and Class Conditions')
P('The colour and shape thresholds for the three fruit classes are all '
  'defined in the HSV_RANGES dictionary and a few module-level constants '
  'in fruit_detector.py. The HSV convention used by OpenCV is: H ∈ [0, '
  '179], and S, V ∈ [0, 255]. The HSV thresholds and the geometric '
  'conditions are summarised in Table 1 and Table 2.')
# Table 1
P('Table 1: HSV colour thresholds for the three fruits.', italic=True)
tbl1 = doc.add_table(rows=4, cols=5)
tbl1.style = 'Light Grid Accent 1'
hdr = tbl1.rows[0].cells
for i, h in enumerate(['Class', 'H', 'S', 'V', 'Note']):
    hdr[i].text = h
rows = [
    ('Banana', '14–40', '80–255', '60–255',
     'Hue shifts to ~15 under warm light, so the lower bound is loosened.'),
    ('Tomato', '0–10 ∪ 170–180', '80–255', '40–255',
     'Red crosses the seam, two segments merged with bitwise_or.'),
    ('Strawberry', '0–10 ∪ 170–180', '30–255', '15–255',
     'Shares the red range with tomato.'),
]
for i, row in enumerate(rows, 1):
    cells = tbl1.rows[i].cells
    for j, v in enumerate(row):
        cells[j].text = v

BLANK()
P('Table 2: Geometric and shape conditions for the three fruits.', italic=True)
tbl2 = doc.add_table(rows=7, cols=4)
tbl2.style = 'Light Grid Accent 1'
hdr = tbl2.rows[0].cells
for i, h in enumerate(['Condition', 'Banana', 'Tomato', 'Strawberry']):
    hdr[i].text = h
rows = [
    ('Min area (px)', '5 000', '1 500', '1 500'),
    ('Max area (px)', '80 000', '60 000', '40 000'),
    ('Aspect ratio', '≥ 1.5 (minAreaRect)', '≤ 1.8 (axis-aligned bbox)', '0.3 ≤ taper ≤ 3.5'),
    ('Circularity (4πA/P²)', '—', '≥ 0.4', '< 0.4 used as shape signal'),
    ('Nearby green', '—', 'green ratio in band above bbox ≤ 5 %',
     'green ratio > 5 % OR shape signal'),
    ('Centre point', 'bbox geometric centre', 'first-moment centroid (cv2.moments)',
     'centre of largest inscribed circle (cv2.distanceTransform)'),
]
for i, row in enumerate(rows, 1):
    cells = tbl2.rows[i].cells
    for j, v in enumerate(row):
        cells[j].text = v

BLANK()
P('The classification of bananas relies mainly on the long, narrow and '
  'yellow shape feature. The minimum area was first set to 800 px, but in '
  'real tests yellowed strawberry calyx leaves (about 1k–3k px and '
  'long-shaped) were often classified as small bananas. The threshold was '
  'therefore raised to 5 000 px, which keeps a four-times safety margin '
  'against real bananas (at least 20k px). The classification of tomatoes '
  'is mainly built on round shape and no nearby green leaves. The '
  'circularity threshold was loosened from the original 0.7 to 0.4 to '
  'allow specular highlights that bite into the contour and reduce '
  'circularity. At the same time, an upper bound of 1.8 on the '
  'axis-aligned bounding-box aspect ratio was added to reject elongated '
  'red objects such as red chillies. The classification of strawberries '
  'is more complex because the calyx may face the camera or be hidden on '
  'the side. The detector therefore uses a "green OR irregular shape" '
  'double-positive condition. It also uses cv2.distanceTransform to find '
  'the centre of the largest inscribed circle as the gripping target, so '
  'the gripper aims at the strawberry\'s wide body rather than the narrow '
  'tip. In addition, the detector computes calyx_dir_base_unit, which is '
  'the XY unit vector from the red centroid to the green calyx centroid, '
  'and passes it to the upper picker to fine-tune the grip offset.')
P('To prevent a single contour from being labelled as more than one '
  'class, the three classes are mutually exclusive on both the shape '
  'axis and the colour axis: the tomato requires circ ≥ 0.4, while the '
  'strawberry\'s shape signal fires when circ < 0.4; the tomato '
  'requires no green nearby, while the strawberry requires green nearby '
  'or an irregular shape. The processing order is set as banana → '
  'tomato → strawberry, but because each detector runs on its own '
  'independent mask, the order only affects the reproducibility of the '
  'output list and does not change the classification result. Every '
  'candidate must finally pass the global threshold CONFIDENCE_MIN = '
  '0.35 to be included in the detection output.')

H3('5.1.3 Camera Calibration Procedure')
P('The accuracy of the vision module depends heavily on camera '
  'calibration, whose purpose is to build a precise mapping between '
  'D415 pixels and robot base-frame 3D coordinates. The calibration '
  'uses an 8×6 chessboard (i.e. 7×5 inner corners, with 30 mm per '
  'square) as the calibration target, and divides the procedure into '
  'three phases.')
P('Phase 1 (origin registration): the operator manually jogs the '
  'robot arm and aligns the gripper tip with the top-left inner corner '
  'of the chessboard. After pressing the confirmation key, the system '
  'records the current TCP position as chess_origin_in_base_m. To '
  'avoid the systematic bias caused by assuming that the chessboard is '
  'strictly aligned with the base axes, it is recommended to also run '
  'touch_three_corners.py and touch the top-right (TR) and bottom-left '
  '(BL) corners. This allows solvePnP to estimate the extrinsics on '
  'the real board direction. All touch results are cached in '
  'logs/last_touched_tcp.json and logs/chess_touched_corners.json to '
  'support the --no-touch reuse mode, so the operator does not have to '
  'jog again every time.')
P('Phase 2 (image capture from the survey pose): the system reads the '
  'pre-taught survey1 joint angles from teach_points.json and calls '
  'slow_move_to_joints to move slowly to that pose, so that the camera '
  'extrinsics used during later detection match the ones solved here. '
  'To avoid sampling before the D415 stream is stable, the pipeline '
  'enters a warm-up phase and uses mean(color) > 5 as the condition '
  'for a valid frame. The system then captures 5 frames in a row and '
  'takes the pixel-wise median as the final captured frame, which '
  'suppresses random noise. The frame is also saved as '
  'logs/calibration_latest.png for later debugging.')
P('Phase 3 (parameter solving): cv2.findChessboardCorners is called '
  'with the CALIB_CB_ADAPTIVE_THRESH + CALIB_CB_NORMALIZE_IMAGE flags '
  'to detect the 35 inner corners, and cv2.cornerSubPix (5×5 window, '
  'up to 30 iterations, convergence threshold 0.01) is used to refine '
  'them at sub-pixel level. Then homography_solver.solve_homography '
  'solves the homography from pixels to the chessboard plane and '
  'returns the reprojection RMS. camera_height_from_homography '
  'further uses the Jacobian of H near the principal point to '
  'estimate the camera height above the table. The error is about 2 % '
  'under a nadir view and about 15 % within a 30° tilt.')
P('The extrinsic estimation is handled by '
  'calibrate_extrinsics.solve_survey1_extrinsics, which calls '
  'cv2.solvePnPGeneric with the SOLVEPNP_IPPE flag. IPPE is a solver '
  'designed for planar targets and returns two candidate solutions at '
  'the same time. The system picks the one whose camera centre z is '
  'above the chessboard plane and whose RMS is lower. Because a 7×5 '
  'chessboard has a direction ambiguity, OpenCV does not always start '
  'the corner indexing from the corner the user touched. The system '
  'therefore runs PnP for four (flip_i, flip_j) combinations '
  '(original, flipped columns, flipped rows, 180° rotation), and '
  'takes the best solution that passes all the gates listed in '
  'Table 3.')
P('Table 3: RMS and physical sanity gates in the calibration procedure.', italic=True)
tbl3 = doc.add_table(rows=5, cols=3)
tbl3.style = 'Light Grid Accent 1'
hdr = tbl3.rows[0].cells
for i, h in enumerate(['Gate', 'Threshold', 'Action when not passed']):
    hdr[i].text = h
rows = [
    ('Homography reprojection RMS', '< 2.0 px',
     'Warning, but still write the calibration file (suggest re-jog)'),
    ('solvePnP reprojection RMS', '< 5.0 px', 'Raise RuntimeError and refuse to write'),
    ('Camera centre z coordinate', '> chess_origin_z',
     'Reject (the IPPE candidate must be above the board plane)'),
    ('Distance from camera centre to survey1 TCP', '< 0.25 m',
     'Reject (physical limit of the wrist-mounted camera)'),
]
for i, row in enumerate(rows, 1):
    cells = tbl3.rows[i].cells
    for j, v in enumerate(row):
        cells[j].text = v

BLANK()
P('After all the gates are passed, run_calibration_core serialises '
  'the origin base coordinate, the homography, the survey joint '
  'angles, the chessboard parameters, the D415 intrinsics, the '
  'reprojection RMS, the camera height above the table, the image '
  'size, and the survey1 extrinsics into session_cal.json, ready to '
  'be loaded at runtime.')

H3('5.1.4 Runtime Calibration Residual Monitoring')
P('Calibration is not "valid forever once it is done". During long '
  'working sessions, the robot arm may drift due to vibration, table '
  'movement or thermal effects in the D415. We therefore introduce a '
  'runtime sanity-check mechanism in survey_capture.py, which '
  'verifies the calibration quality before each detection. Every '
  'time capture_fruits is called for a detection, it first moves the '
  'arm to the survey1 pose (the same as during calibration), '
  'captures the RGB and depth images, and then calls '
  'cv2.findChessboardCorners again. The detected corners are '
  'projected back to the chessboard plane through '
  'session_cal.h_pixel_to_chess_mm and compared with the nominal '
  '30 mm grid to compute the residual in millimetres. The system '
  'then decides what to do based on the residual size, as shown in '
  'Table 4.')
P('Table 4: Runtime residual levels and actions.', italic=True)
tbl4 = doc.add_table(rows=4, cols=3)
tbl4.style = 'Light Grid Accent 1'
hdr = tbl4.rows[0].cells
for i, h in enumerate(['Residual', 'Status colour', 'System behaviour']):
    hdr[i].text = h
rows = [
    ('< 3 mm', 'Green', 'Normal, continue with detection'),
    ('3 mm ≤ residual < 10 mm', 'Amber',
     'Add to warnings list, HUD shows a warning string'),
    ('≥ 10 mm', 'Red', 'Raise RuntimeError and force a re-calibration'),
]
for i, row in enumerate(rows, 1):
    cells = tbl4.rows[i].cells
    for j, v in enumerate(row):
        cells[j].text = v
BLANK()
P('This design changes calibration-quality monitoring from a one-time '
  'check at start-up into a self-check on every frame, which clearly '
  'improves the long-term stability of the system.')

H3('5.1.5 Object Detection Visualisation')
P('The drawing of the detection boxes is intentionally separated '
  'from fruit_detector.py into the _annotate() function in '
  'picker_viewer.py. This design keeps the detector clean — it only '
  'returns the geometric and classification results and does not '
  'pollute the original frame — and makes the visualisation an '
  'independent overlay that can be turned on or off freely. The '
  'Detection object returned by the detector already carries all the '
  'data needed for drawing, including bbox, center_px, confidence, '
  'and fruit_type. To make the three fruit classes easy to tell '
  'apart, each one has its own BGR colour: banana yellow '
  '(0, 255, 255), tomato red (0, 0, 255), and strawberry pink-purple '
  '(200, 100, 255). The HUD text is white and warning overlays are '
  'amber. The drawing core is a single for-loop. For each detection, '
  'three drawing operations are performed: cv2.rectangle draws the '
  'axis-aligned box; cv2.putText writes the class and confidence '
  'label at the top-left of the box; and a filled circle marks the '
  'gripping target at center_px.')

H2('5.2 Autonomous Mode')
P('The autonomous mode is implemented as a finite state machine '
  '(FSM) inside sorting_controller.py. The FSM cycles through '
  'GO_HOME, APPROACH, DESCEND, CLOSE_GRIPPER, ASCEND_PICK, '
  'MOVE_TO_BASKET, DESCEND_PLACE, OPEN_GRIPPER, ASCEND_PLACE and '
  'DONE for each fruit. The picker dispatcher in picker_viewer.py '
  'orders the queue as bananas → tomatoes → strawberries, and '
  'within each class the order follows the detector\'s output list.')
P('On top of the geometric centroid returned by the detector, the '
  'picker applies type-specific empirical biases that were tuned in '
  'lab. For the strawberry path, the grasp target is shifted by '
  '+2 cm in the calyx direction (the unit vector from the body '
  'centroid to the green-pixel centroid) and by −3 cm in the global '
  'x axis, partially cancelling the controller\'s +5 cm global '
  'x-bias to give a net effective +2 cm. For the tomato path, the '
  'grasp target is shifted by +2 cm in x, giving a net effective '
  '+7 cm. Banana picks use the global bias only. These biases were '
  'tuned by running single-fruit grasps in lab and adjusting the '
  'offsets until the gripper consistently landed on the body of the '
  'fruit. The full tuning trace is documented as inline comments at '
  'picker_viewer.py:178–208.')
P('The FSM also emits a structured trace log on every transition, '
  'including PICK_ATTEMPT entries with the target Cartesian pose '
  'and confidence, and GRIPPER_READBACK entries with the commanded '
  'and actual gripper closure. These trace events are the source of '
  'the stall data analysed in Section 7.6.')

H2('5.3 Remote-Controlled Mode')
P('The remote-controlled mode is implemented as a Tkinter GUI in '
  'python/main_gui.py. The interface follows a two-column desktop '
  'layout. The left column hosts a real-time video feed from the '
  'D415 camera at 640×360 pixels, while the right column groups the '
  'controls by function into several panels covering emergency '
  'stop, mode switching, manual teleoperation, gripper control, '
  'camera operations, and automatic sorting.')

H3('5.3.1 Layout and Mode Switching')
P('The main window, titled "QArm Fruit Sorting — GUI", is '
  'constructed with ttk.Frame and the grid geometry manager into a '
  'two-column layout. One subtle pitfall deserves explicit '
  'mention: when a Label widget is initialized without an image '
  'attribute, its width and height parameters are interpreted in '
  'character units rather than pixels, so a value of 640 '
  'corresponds to roughly 3 840 pixels and is enough to push the '
  'right column off the screen. To prevent this, the _build_ui '
  'method preallocates a 640×360 black PhotoImage and assigns it '
  'as the initial image. The right column is divided from top to '
  'bottom by ttk.Separator instances into five functional regions: '
  'the emergency stop button, mode switching, manual teleoperation '
  '(covering XYZ jogging, basket positioning, gripper, and camera '
  'operations), the automatic sorting button group, and the status '
  'bar at the bottom.')
P('Mode switching is implemented as a binary selection through '
  'ttk.Radiobutton widgets, bound to self.mode_var. When the user '
  'changes mode, _on_mode_change() iterates through the lists '
  'self.manual_buttons and self.auto_buttons and sets the state of '
  'every button in the active group to normal while disabling the '
  'other group. This eliminates command conflicts that would '
  'otherwise arise if the operator pressed manual teleoperation '
  'keys during automatic batch execution.')

H3('5.3.2 Emergency Stop')
P('The emergency stop button sits at the top of the right column '
  'and carries the highest visual weight in the entire interface. '
  'It is implemented as a tk.Button labelled EMERGENCY STOP, with '
  'a high-contrast red background, white foreground, bold Arial '
  '14 pt typography, and binds to _emergency_stop. To address the '
  'trade-off in interrupt granularity inherent to real-time '
  'control systems, the emergency stop is realised through two '
  'layers of interruption. The first layer is a flag-based '
  'interrupt: a call to self.abort_event.set() raises the '
  'threading.Event so that every long-running worker thread, '
  'including teleoperation, basket motion, survey1, and the '
  'automatic mode, checks self.abort_event.is_set() at every loop '
  'iteration and returns immediately. The second layer is a '
  'pose-freezing burst: because in-flight motion commands may '
  'still be propagating through the 200-step interpolation loop '
  'in slow_move_to_joints, the code reads the current joint '
  'angles and gripper state and writes them back to the driver '
  'repeatedly at ESTOP_FREEZE_REPEAT = 30 iterations and '
  'ESTOP_FREEZE_INTERVAL = 0.01 s, totalling roughly 300 ms.')

H3('5.3.3 Manual Teleoperation')
P('The manual teleoperation region is partitioned along physical-'
  'space dimensions into four subregions, responsible for '
  'Cartesian jogging, basket positioning, gripper control, and '
  'camera operations. Cartesian XYZ jogging is realised by six '
  'buttons arranged in a 3×4 grid: Y+ and Z+ occupy row 0, X− and '
  'X+ occupy row 1, and Y− and Z− occupy row 2. Each press issues '
  'a displacement command of TELEOP_STEP_M = 0.02 m (i.e. 2 cm), '
  'with a fixed motion duration of TELEOP_SECONDS = 0.6 s and an '
  'interpolation count of TELEOP_STEPS = 40. The callback '
  '_teleop(dx, dy, dz) executes the following sequence: it clears '
  'abort_event, reads the current joint angles and gripper state '
  'from the driver, calls forward_kinematics to obtain the '
  'current end-effector position ee_pos, adds the increment '
  '(dx, dy, dz) to obtain the target new_pos, and then invokes '
  'inverse_kinematics(new_pos, gamma=0.0). If the IK solver '
  'fails, the error is reported on the status bar and the '
  'function returns. Otherwise, _safe_move_to_joints interpolates '
  'the joints to the target pose, and the emergency stop is '
  'allowed to interrupt the motion at any point. As shown later '
  'in Section 7.5, the gamma=0.0 constraint is one of the two '
  'control-architecture defects that make teleop unusable in '
  'practice, because the IK redistribution under a fixed '
  'orientation constraint causes the end-effector to drift '
  'downward in Z on every Cartesian step.')
P('Basket positioning is provided by a row of three buttons for '
  'Strawberry, Tomato, and Banana, each bound to '
  '_move_to_basket(fruit_type). The callback first reads the '
  'preset basket pose from controller.BASKETS[fruit_type], then '
  'performs IK, and finally calls _safe_move_to_joints with a '
  'fixed duration of BASKET_SECONDS = 2.0 s. Gripper control '
  'consists of two buttons, Open and Close, bound to '
  '_gripper_open and _gripper_close respectively. These methods '
  'invoke controller.set_gripper_ramp with GRIP_OPEN or '
  'GRIP_CLOSE. set_gripper_ramp issues a ramped command, which '
  'prevents the gripper from snapping abruptly to the endpoint '
  'and damaging the target object. Its return value, named '
  'actual, denotes the position actually held and is reported on '
  'the status bar. The camera subregion contains two buttons. '
  'The Goto Survey1 button moves the manipulator to the imaging '
  'pose self.session_cal.survey_pose_joints_rad with a motion '
  'duration of SURVEY_SECONDS = 2.5 s. The Refresh detect button '
  'does not move the arm; instead, it directly invokes the '
  'hooked capture_fruits to perform the full sequence of image '
  'capture, chessboard self-check, and fruit detection.')

H3('5.3.4 Automatic Sorting Buttons and Status Bar')
P('The automatic sorting region contains four buttons whose '
  'enabled and disabled states are managed jointly through the '
  'self.auto_buttons list. The ALL (B → T → S) button picks every '
  'banana, tomato, and strawberry in order and binds to '
  '_auto_all, while the remaining three buttons (Bananas only, '
  'Tomatoes only, Strawberries only) target a single category and '
  'bind to _auto_one(fruit_type). Both callbacks first invoke '
  '_install_abort_observer() to register '
  'controller.tick_observer = _abort_observer. The observer '
  'checks abort_event at every FSM tick; once the flag is set, '
  'it forces controller.state to State.DONE and raises '
  'RuntimeError. The execution body itself reuses '
  'picker_viewer._pick_category, which allows main_gui and '
  'picker_viewer to share the same implementation at the '
  'automatic batch layer and avoids duplicate code.')
P('The status bar resides at the bottom of the right column, '
  'implemented as a ttk.Label bound to a tk.StringVar named '
  'self.status_var, configured with wraplength = 260 to enable '
  'line wrapping, and styled in dark green. Every long-running '
  'operation, whether successful or failed, updates the bar '
  'through the _set_status(msg) method. Internally this method '
  'uses self.root.after(0, lambda: …) to marshal the string onto '
  'the main thread, which avoids the "main thread is not in '
  'main loop" exception that would otherwise arise from '
  'background threads operating directly on Tk widgets.')

# --- 6. Soft Gripper ---
PAGEBREAK()
H1('6. Soft Gripper Design and Analysis')

H2('6.1 Design Selection')
P('The gripper design adopted in this project is based on the '
  'Fin Ray Effect, a biomimetic compliant mechanism in which '
  'structural ribs deflect toward an applied contact force. This '
  'behaviour arises from the internal cross-rib architecture, '
  'which redistributes compressive loads across the finger '
  'structure rather than concentrating stress at a single contact '
  'point, enabling passive shape conformance to irregular object '
  'geometries without active sensing.')
P('Two alternative soft gripper paradigms were evaluated prior '
  'to this selection. PneuNet actuators were excluded because '
  'the compressed-air supply infrastructure they require is '
  'incompatible with the QArm\'s existing Dynamixel XC430 '
  'servo-driven hardware, introducing unnecessary system '
  'complexity. Granular jamming, while effective for highly '
  'irregular geometries, does not admit closed-form mechanical '
  'analysis, making rigorous safety-factor verification '
  'impractical within the scope of this coursework. The Fin Ray '
  'topology, by contrast, is analytically tractable under '
  'Euler–Bernoulli beam bending theory, integrates directly with '
  'tendon actuation via the XC430 servo, and requires no '
  'additional hardware beyond the existing manipulator platform.')

H2('6.2 Geometry and Materials')
P('The gripper comprises two symmetrical Fin Ray fingers, each '
  '80 mm in length and 35 mm in width, a footprint determined by '
  'the largest target object, the tomato. Structural ribs are '
  'fabricated from TPU 85A via fused deposition modelling. With '
  'an elastic modulus of approximately 20 MPa, this material '
  'provides sufficient rigidity during grasp initiation while '
  'retaining the elastic compliance needed for sustained '
  'contact. Each rib has a rectangular cross-section with '
  'thickness graded from 1.5 mm at the fingertip to 2.0 mm at '
  'the base; this taper intentionally concentrates deflection '
  'toward the distal end, where conformance to object curvature '
  'is most critical. Rib lengths range from 15 mm at the tip to '
  '30 mm at the base: shorter ribs near the contact zone limit '
  'tip deflection, while longer basal ribs provide the '
  'structural lever arm required for force transmission. The '
  'contact surfaces are lined with Silicone A30, a soft, '
  'food-safe elastomer with a conservatively estimated friction '
  'coefficient of μ = 0.6, ensuring reliable grip across the '
  'range of fruit surface textures encountered in this task. '
  'Total gripper mass is estimated at 60 to 80 g, remaining '
  'within the 100 g end-effector payload budget and preserving '
  'the effective working payload of the QArm.')

H2('6.3 Mechanical Analysis')
P('Mechanical performance was evaluated through four analytical '
  'models, each addressing a distinct failure or functional '
  'criterion. The no-slip condition requires '
  'F_grip ≥ m·g / (n·μ), where m is object mass, g gravitational '
  'acceleration, n the number of contact fingers, and μ the '
  'friction coefficient. With μ = 0.6 and n = 2, this '
  'establishes the minimum gripping force for each target '
  'object, verified in the safety factor analysis below. '
  'Contact stress at the finger–object interface is estimated '
  'using the Hertzian contact model, giving a peak pressure of '
  'P_max = 3F / (2πa²), where a is the effective contact patch '
  'radius; this provides an upper bound on surface stress '
  'against which the bruising threshold of each fruit can be '
  'checked. Rib deflection under load is modelled as a '
  'cantilever beam under Euler–Bernoulli assumptions, yielding '
  'tip deflection δ = FL³ / (3EI) and maximum bending stress '
  'σ_max = FL / (bh²/6), where L is rib length, E the elastic '
  'modulus of TPU 85A, and I = bh³/12 the second moment of '
  'area of the rectangular cross-section. Available tendon '
  'force is derived from the Dynamixel XC430 torque '
  'specification: adopting a conservative operating torque of '
  'τ = 0.5 N·m and a pulley radius of r = 5 mm gives a tendon '
  'tension of F = τ/r = 100 N, which comfortably exceeds the '
  'gripping force requirement across all target objects.')
P('Two candidate rib thicknesses were evaluated under a '
  'representative load of F = 1 N. The 1.5 mm rib gives '
  'I = 5.625 × 10⁻¹³ m⁴, a tip deflection of δ = 0.46 mm, and a '
  'peak bending stress of σ_max = 33.3 MPa. While the '
  'deflection is mechanically permissible, the bending stress '
  'approaches the tensile strength of TPU 85A at roughly '
  '30–40 MPa, leaving a low structural margin; this '
  'configuration was therefore judged functional but marginal. '
  'Increasing rib thickness to 2.0 mm reduces tip deflection '
  'to δ = 0.20 mm and peak bending stress to '
  'σ_max = 18.75 MPa, improving the stress margin at the cost '
  'of reduced finger compliance. The 2.0 mm thickness was '
  'selected as the baseline design, accepting the moderate '
  'reduction in passive conformability in exchange for greater '
  'structural reliability over repeated grasping cycles.')

H2('6.4 Safety Factor Analysis')
P('Safety factor was defined as SF = F_bruising / F_min, where '
  'F_bruising is the empirically established bruising threshold '
  'and F_min the minimum gripping force required by the '
  'no-slip condition. The banana yields F_min = 1.23 N against '
  'F_bruising = 10 N, giving SF = 8.1, reflecting its '
  'relatively firm peel and low density. The tomato gives '
  'F_min = 1.64 N and F_bruising = 5 N, with SF = 3.0; the '
  'lower margin is consistent with its thin, pressurised skin, '
  'which is more susceptible to localised stress '
  'concentrations. The strawberry package is the most '
  'demanding case: the aggregate mass of six strawberries '
  'raises F_min to 3.27 N, while the packaging material '
  'bruising threshold of 8 N gives SF = 2.4. This minimum '
  'safety factor holds only when Force-Sensitive Resistor '
  'feedback is active, confirming that closed-loop gripping '
  'force regulation is a necessary condition for safe '
  'strawberry-package handling rather than an optional '
  'enhancement. Across all three object types, safety factors '
  'exceed unity by a sufficient margin, validating that the '
  'gripper can complete the pick-and-place task without '
  'inducing fruit damage under the specified operating '
  'conditions.')

# ====================================================================
# PART 2 — VALIDATION (7 pages)
# ====================================================================
PAGEBREAK()
H1('7. Experimental Validation and Results')

H2('7.1 Test Methodology')
P('The project did not include a Simulink dynamic simulation '
  'of the QArm hardware, so a conventional sim-to-real '
  'comparison was not possible. To still characterise the '
  'system properly, we designed a layered testing strategy on '
  'the physical robot. Four test modes were used. The first '
  'mode was a single-fruit fixed-position pre-trial, in which '
  'each fruit type was placed alone at six workspace points '
  '(P1 to P6) one at a time. The second and third modes were '
  'two autonomous repeatability runs, each with a full set of '
  '14 fruits in standard upright placement. The fourth was an '
  'autonomous robustness run that intentionally introduced '
  'adversarial placements such as adjacent, inverted, and '
  'rotated fruits. Finally, a teleoperation test was carried '
  'out by two human operators. These four modes together '
  'cover both the best case and the worst case of the system, '
  'and they let us isolate where each kind of failure '
  'actually comes from.')
P('The hardware and software setup was as follows. The '
  'platform was a Quanser QArm 4-DOF manipulator with a rigid '
  'two-finger gripper, paired with an Intel RealSense D415 '
  'stereo-IR depth camera mounted overhead. A 7×5 inner-corner '
  'chessboard (8×6 squares at 30 mm) served as the '
  'calibration fiducial. The protocol for each autonomous run '
  'started with running session_cal.py with a chessboard '
  'touch-probe to align the camera frame to the QArm base '
  'frame, after which the fruits were placed by hand on the '
  'workspace and photographed for the ground-truth record. '
  'The picker was then started in ALL (B → T → S) mode. '
  'During the run we logged every detection frame, every '
  'PICK_ATTEMPT event, the commanded and read-back gripper '
  'closures, and the joint angles at grip and release. After '
  'the run, two team members independently checked the '
  'photographs against the basket contents and the log, and '
  'any discrepancy between the log and the visual record was '
  'resolved in favour of the visual record. This protocol was '
  'important because we found early in testing that the '
  'controller reports a successful "Placed" event even when '
  'the gripper has closed empty: the controller cannot '
  'directly tell whether anything is between the fingers, '
  'only that the gripper command was issued and the '
  'lift-and-release motion completed. The visual-record check '
  'is therefore the single source of truth for grasp outcomes '
  'in this section. This is also why the per-attempt grasp '
  'success rate (33/46 = 71.7 %) is reported separately from '
  'the per-fruit task completion rate (33/41 = 80.5 %): the '
  'same fruit can generate multiple attempts in the log, and '
  'the two metrics tell different stories about the system.')
P('Classification accuracy is the number of fruits whose '
  'category was correctly identified by the detector, divided '
  'by the number of fruits in the workspace. Grasp success '
  'rate is reported in the two ways described above '
  '(per-attempt and per-fruit). Basket placement accuracy is '
  'the proportion of picked fruits that were dropped into the '
  'correct basket. Task completion rate is the end-to-end '
  'metric, defined as the number of fruits both picked and '
  'placed in the correct basket, divided by the total number '
  'of fruits in the workspace. Cycle time is the average '
  'duration of one pick-and-place loop. We also recorded the '
  'gripper stall value from each pick attempt, defined as the '
  'difference between the commanded gripper closure and the '
  'actual readback. As shown in Section 7.6, this single '
  'value turned out to be a useful diagnostic signal for '
  'separating different failure modes.')

H2('7.2 Pre-Trial: Single-Fruit Fixed-Position Performance')
P('The pre-trial was designed to isolate the lowest-level '
  'question: given a single fruit placed at a known location, '
  'can the system pick it up? This removes multi-target '
  'scheduling, adjacency effects, and visual clutter, so any '
  'failure here points to a fundamental capability problem '
  'rather than a planning issue.')
P('The results are summarised in Figure 1, which mirrors the '
  'actual 2-by-3 spatial layout of the workspace. Banana '
  'achieved 6 out of 6, although it should be noted that we '
  'used 3 physical bananas, each spanning two adjacent '
  'positions because of their elongated shape; this still '
  'gives full coverage of the six positions but the effective '
  'sample size is 3 rather than 6. Strawberry achieved 5 out '
  'of 6: the only miss was at P2, where the strawberry was '
  'misclassified as a tomato by the detector and was '
  'therefore skipped by the "strawberries-only" picker. '
  'Position P5 needed two attempts; the first one missed and '
  'the second one succeeded. Tomato performed clearly the '
  'worst, with only 2 out of 6 positions yielding a '
  'successful pickup (P2 and P3). On the other four positions '
  'the gripper either closed past the fruit completely '
  '(P4, P5, P6, where stall values fell to 0.02–0.023) or '
  'held it briefly and dropped it during the lift phase (P1, '
  'where stall was 0.192 — within the normal grip range, yet '
  'the fruit still slipped). The overall first-attempt '
  'success rate of the pre-trial was 13 out of 18 (72.2 %).')
FIG('figC_pretrial_heatmap.png',
    'Figure 1: Pre-trial single-fruit pickup outcome by '
    'workspace position and fruit type.', width_in=5.0)
P('The pre-trial therefore exposed the tomato as the weakest '
  'link in the system. Because tomatoes are smooth, fairly '
  'heavy, and roughly spherical, the rigid two-finger gripper '
  'does not have enough friction to hold them reliably, and '
  'small errors in the visual centroid translate into large '
  'errors in the actual contact geometry. The P1 case is '
  'particularly informative because it shows that even a '
  'stall value in the normal range does not guarantee a '
  'successful pickup — a finding we return to in Section 7.6. '
  'This finding directly motivates the soft gripper design '
  'described in Chapter 6.')

H2('7.3 Autonomous Mode: Repeatability (Trial 1 and Trial 2)')
P('The two repeatability runs each used the same workspace '
  'layout: 3 bananas, 5 tomatoes, and 6 strawberries placed '
  'in standard upright orientation, with non-overlapping '
  'spacing. The picker ran the full ALL (B → T → S) sequence, '
  'meaning it cleared all bananas first, then all tomatoes, '
  'then all strawberries.')
P('Trial 1 finished in 6 minutes 35 seconds and reached a '
  'task completion rate of 13 out of 14 (92.9 %). The single '
  'failure was F09, a strawberry: the detector identified it '
  'correctly with confidence 1.00, the picker commanded the '
  'close, the gripper closed to a stall of only 0.0226 (the '
  'lowest stall of any strawberry attempt in the trial, well '
  'below the typical successful-grip range of 0.05–0.08 for '
  'strawberries), the lift-and-place sequence completed, and '
  'the controller logged a successful pick. Visual '
  'inspection confirmed that the gripper had in fact closed '
  'without retaining the fruit. This is exactly the kind of '
  'false-positive log entry that motivated the visual-'
  'verification protocol described in Section 7.1, and it is '
  'reported under the per-fruit task completion rate (13/14) '
  'rather than under the log-derived per-attempt count.')
P('Trial 2 had F08 missing from the workspace because of a '
  'placement mistake, so we report it as 13 out of 13 '
  '(100 %) and finished in 8 minutes 57 seconds. Across both '
  'runs, basket placement accuracy was 100 %: every fruit '
  'that was successfully picked up was dropped into the '
  'correct basket, with no misclassification at the '
  'placement stage.')
FIG('figA_autonomous_completion.png',
    'Figure 2: Autonomous task completion rate across the '
    'three trials, broken down by fruit type.', width_in=5.5)
P('The most encouraging part of the result is the cycle '
  'time. The average time per pick attempt was 28.2 seconds '
  'in Trial 1 and 29.8 seconds in Trial 2, with very small '
  'spread between attempts (the per-attempt distribution was '
  'tight in the 27.7 to 33.8 second range). This shows that '
  'the kinematics, motion planning, and gripper control loop '
  'are highly repeatable. The longer total time of Trial 2 '
  'was caused by the picker spending more attempts on '
  'tomatoes (some of them needed re-grasping), not by the '
  'cycle becoming unstable. Specifically, Trial 2 logged 9 '
  'tomato PICK_ATTEMPT events on 5 tomato fruits — of those '
  '9 attempts, 3 had stall in the normal grip range '
  '(0.18–0.22) and successfully retrieved a tomato, 4 had '
  'stall around 0.02 (empty grasp), and 2 had marginal stall '
  '(0.05–0.07) on retries that eventually succeeded.')
P('The cycle time itself can be decomposed qualitatively '
  'into four phases: detection (a single overhead capture '
  'and depth fusion), pre-grasp move from the survey pose to '
  'the fruit, grasp closure with stall confirmation, and '
  'lift-and-place from the fruit to the correct basket. '
  'Among these, the place leg is by far the longest because '
  'it traverses the largest joint distance, which is fixed '
  'by the workspace geometry and by the conservative '
  'slow_move_to_joints duration used for safety. This means '
  'cycle time is dominated by the kinematic constants and '
  'the safety-tuned motion durations, not by the detection '
  'or planning layers, and is therefore a property of the '
  'platform rather than of the algorithm. Considering that '
  'this baseline performance covers detection, depth '
  'fusion, IK, gripper control, and basket sorting in a '
  'single closed loop running on the physical QArm, the '
  'autonomous baseline is a strong outcome for the project.')

H2('7.4 Autonomous Mode: Robustness (Trial 3)')
P('Trial 3 deliberately introduced placements that we '
  'expected to challenge the system, in order to map out '
  'where the failure boundaries actually are. The same 14 '
  'fruits were used, with the following modifications: F03 '
  '(banana) and F04 (tomato) were placed touching each '
  'other; F05, F06, and F08 were inverted tomatoes; F09 and '
  'F11 were inverted strawberries; F13 and F14 were '
  'strawberries rotated more than 90 degrees from the '
  'canonical pose. The remaining fruits were placed '
  'normally.')
P('The run finished in 6 minutes 56 seconds and achieved a '
  'task completion rate of 7 out of 14 (50 %). The failures '
  'split into three categories.')
P('The first category is adjacent-fruit interference (F03 '
  'and F04). The detector did identify F03 as a banana, '
  'but the live detection frames captured at this moment '
  'show that F04 — the adjacent tomato — was confidently '
  'misclassified as a strawberry, with reported confidence '
  '0.86 in one frame and 1.00 in a slightly later frame. '
  'Inspection of the detector source code explains this '
  'directly: the strawberry path '
  '(fruit_detector.py:320–321) accepts a red blob if either '
  'a green calyx signal exceeds 0.05 OR the contour '
  'circularity drops below 0.4. When the tomato is touching '
  'a banana, the banana\'s stem region or the contact '
  'shadow can leak green pixels into the bbox neighbourhood '
  '(the calyx signal is searched in the bbox plus 25 pixels '
  'of padding, fruit_detector.py:357–358), and the '
  'disturbed contour can also fall below the circularity '
  'threshold. Either pathway is sufficient to flip the '
  'label from tomato to strawberry. Because the picker was '
  'in the "banana phase" at this moment, the tomato had '
  'been misclassified into the strawberry queue rather than '
  'the banana queue, and the picker pursued the actual '
  'banana — but the gripper hit both fruits at once. The '
  'recorded stall values confirm this: banana attempt #3 in '
  'the Trial 3 log returned stall = 0.4300 and attempt #4 '
  'returned stall = 0.4727, meaning the gripper closed '
  'only to roughly 53–57 % of the commanded position. '
  'Normal banana stall in this run was 0.18–0.22, so a '
  'stall above 0.4 is a clear signature of mechanical '
  'jamming on a hard object — consistent with the gripper '
  'being stuck on the harder, smoother tomato wedged '
  'against the banana.')
P('The second category is inverted-pose failure (F08, F09, '
  'F11). For inverted tomatoes (F08), the geometric '
  'centroid and the depth fusion together gave a target '
  'point that was offset from the actual contact surface, '
  'and the gripper closed past the fruit twice before we '
  'withdrew it. For inverted strawberries (F09 and F11), '
  'the failure mechanism is more interesting and is best '
  'explained directly from the picker source. The '
  'strawberry pick logic (picker_viewer.py:178–225) '
  'applies two empirical biases on top of the detected '
  'centroid: a 2 cm offset in the direction of the calyx '
  '(_STRAWBERRY_CALYX_BIAS_M = 0.02), and a global x-axis '
  'offset of −3 cm (_STRAWBERRY_X_BIAS_M = −0.030, '
  'partially cancelling the controller\'s +5 cm global '
  'x-bias to give a net effective +2 cm). The widest-'
  'inscribed-disk centroid algorithm '
  '(fruit_detector.py:277–298) was tuned for upright '
  'strawberries, where the wide body forms the dominant '
  'projection from above and the algorithm correctly '
  'places the centroid inside that body. On an inverted '
  'strawberry the projected outline changes shape: the '
  'centroid algorithm tends to land near the calyx stem '
  'rather than inside the wide fruit body, and the '
  'calyx-direction unit vector then points along the same '
  'axis. The 2 cm bias pushes the grasp another 2 cm in '
  'that direction, which lands the gripper outside the '
  'actual contact surface. Both inverted strawberries '
  '(F09 and F11) produced an empty grasp on every attempt, '
  'with stall values of 0.0206 and 0.0231 respectively in '
  'the Trial 3 log — consistent with the A-type empty-'
  'grasp signature defined in Section 7.6.')
P('The third category is detection rejection (F13 and '
  'F14). With the strawberry rotated more than 90 degrees '
  'from canonical, the calyx is no longer visible from '
  'above, AND the side-profile of the strawberry looks '
  'more circular. Reading the detector code, this causes '
  'both strawberry acceptance paths to fail: the calyx '
  'ratio drops below 0.05 (no green found) and the '
  'circularity rises above 0.4 (failing the shape signal). '
  'The contour is therefore rejected before any Detection '
  'object is produced, which is why the Trial 3 log '
  'contains no PICK_ATTEMPT entry for F13 or F14 — the '
  'picker simply never saw them.')
P('Notably, F05 and F10 are partial successes that show '
  'the system can sometimes handle adversarial poses. F05 '
  'was an inverted tomato that we expected to fail (and '
  'had failed in earlier informal tests), but in Trial 3 '
  'it was picked cleanly on the first attempt with stall '
  '0.1321. F10 was a strawberry rotated about 45 degrees '
  'from canonical, which is below the angle at which the '
  'side-profile becomes circular enough to defeat the '
  'shape signal; it was detected and grasped successfully.')
P('Although a 50 % completion rate looks weak compared to '
  'Trial 1 and Trial 2, it should be read as a stress '
  'test. The drop from a 96 % baseline to 50 % under '
  'adversarial input is informative because it tells us '
  'exactly which conditions break the system, and the '
  'failure modes are structural and traceable rather than '
  'random.')

H2('7.5 Teleoperation Mode')
P('The teleoperation test was carried out by two team '
  'members. The test rule was 6 fruits per operator '
  '(1 banana, 3 tomatoes, 2 strawberries), with 30 '
  'seconds of warm-up time on the keys before recording '
  'started. The original plan included a third operator, '
  'but after watching the first two attempts the third '
  'tester declined further participation, and we chose to '
  'report this as a usability finding rather than as '
  'missing data.')
P('The combined result was 0 out of 12. Neither operator '
  'was able to place a single fruit in the correct basket. '
  'Tester 1 took 8 minutes 57 seconds, recorded 4 grip '
  'failures and 4 collision events, and rated the '
  'experience 1 out of 5. Tester 2 took 4 minutes 47 '
  'seconds, recorded 3 grip failures and 3 collision '
  'events, and likewise rated the experience 1 out of 5. '
  'The tester-2 run was shorter only because the operator '
  'gave up earlier after experiencing the same failures. '
  'Both runs failed in essentially the same way: any '
  'fruit that was picked up gradually slipped out as the '
  'operator tried to move it toward the basket, and fine '
  'adjustments near the pick pose caused the gripper to '
  'crash into the workspace surface.')
P('We investigated the cause directly in the source code. '
  'Two issues were found. First, the GUI labels the '
  'manual jog buttons as "base frame, 2 cm step", and the '
  'function _teleop(dx, dy, dz) in main_gui.py:512–526 '
  'reads the current joint angles, runs forward '
  'kinematics to get the end-effector position, adds the '
  'requested Cartesian step, and then calls inverse '
  'kinematics with gamma = 0.0 to get the new joint '
  'target. The actual motion is then a joint-space '
  'interpolation. Because the IK solution under the fixed '
  'orientation constraint redistributes all four joint '
  'angles on every step, the joint-space path between two '
  'close Cartesian targets does not stay on the '
  'requested straight line, and the end-effector tends to '
  'drift downward in Z. Operators reported exactly this '
  'behaviour: pressing X+ or Y+ caused the gripper to '
  'also sink toward the table.')
P('Second, the gripper command in teleop is anchored to '
  'the readback of the actual gripper position rather '
  'than to the operator\'s last commanded value. At line '
  '515 of main_gui.py, every teleop step reads the '
  'current gripper value from the driver via '
  'driver.read_all() and then passes that value into '
  '_safe_move_to_joints (line 524) as the gripper '
  'setpoint for the upcoming move. If the operator has '
  'closed the gripper on a fruit and the fruit '
  'subsequently slips slightly under load, the actual '
  'gripper opens slightly to match, the readback returns '
  'the opened value, and the next teleop press commands '
  'that opened value. The grip therefore drifts in the '
  'same direction as the slip, rather than being held at '
  'the original closure target. Each subsequent press '
  'reinforces the slip rather than correcting it, and '
  'the fruit eventually falls out before reaching the '
  'basket.')
P('These two control-layer issues fully explain the 0 % '
  'result. The teleop test is therefore best read as a '
  'control group: when the autonomous wrapper is removed, '
  'the underlying low-level layer is not yet usable for '
  'human operators, which makes the autonomous mode\'s '
  '96 % baseline result look more meaningful by '
  'comparison.')

H2('7.6 Failure Mode Analysis')
P('To make the failure pattern across all four test '
  'modes easier to interpret, we extracted the gripper '
  'stall value from every recorded pick attempt and '
  'grouped it by fruit type and outcome. The result is '
  'shown in Figure 3 and reveals three clear classes of '
  'failure with strikingly clean separation in the stall '
  'axis.')
FIG('figB_stall_distribution.png',
    'Figure 3: Distribution of gripper stall values across '
    'all pick attempts, grouped by fruit type and outcome. '
    'Three failure classes (A, B, adjacency) are visible.',
    width_in=5.5)
P('The first class is the empty-grasp failure, which we '
  'call A-type. In these cases the gripper closes almost '
  'completely (stall in the 0.02 range) because there is '
  'nothing inside it. This pattern appears at five tomato '
  'attempts in the pre-trial (T#4 through T#8, stall '
  '0.020–0.023), at the inverted strawberries F09 (stall '
  '0.0206) and F11 (stall 0.0231) in Trial 3, and at the '
  'inverted tomato F08 in Trial 3. The root cause is a '
  'wrong grasp pose, either from the type-specific bias '
  'applied downstream of detection or from the widest-'
  'disk centroid algorithm landing on the calyx stem '
  'rather than inside the body when the strawberry is in '
  'a non-canonical orientation (as discussed in '
  'Section 7.4).')
P('The second class is the slip-during-ascent failure, '
  'which we call B-type. The stall is in the normal range '
  'for a successful grip (around 0.19) but the fruit '
  'slips out during the lift phase. The clearest example '
  'is the tomato P1 in the pre-trial, where stall was '
  '0.192 — well within the typical successful-grip range '
  '— yet the tomato fell out of the gripper during '
  'ascent. This kind of failure cannot be detected from '
  'the stall value alone, because the gripper looks like '
  'it grabbed something. This is an important lesson: '
  'stall is a necessary but not sufficient signal for '
  'grasp success, and any future closed-loop verification '
  'will need a second sensor (for example, a low-cost '
  'weight sensor on the gripper or a downward-looking '
  'camera between the fingers) to confirm retention '
  'through the lift phase.')
P('The third class is the adjacency failure (F03 and F04 '
  'in Trial 3), where the stall is anomalously high '
  '(0.4300 and 0.4727 respectively) because the gripper '
  'is mechanically jammed between two fruits. The exact '
  'stall values are diagnostic: a stall of around 0.20 '
  'is the natural compliance of the rigid finger pads '
  'against soft fruit material, while a stall above 0.4 '
  'means the gripper has stopped against something rigid '
  'enough to prevent further closure. In this case it '
  'was the harder tomato wedged against the softer '
  'banana.')
P('A useful way to read these three classes together is '
  'to ask which layer of the system is responsible for '
  'each. A-type failures originate in the perception-to-'
  'action interface: the detector finds the fruit and '
  'reports a centroid, but the bias logic and centroid '
  'algorithm that translate the centroid into a grasp '
  'pose are not robust to non-canonical orientations. '
  'B-type failures originate in the physical actuation '
  'layer: the grasp pose is roughly correct, the gripper '
  'closes, but the friction and geometry of the rigid '
  'two-finger design cannot retain the fruit through '
  'the lift. Adjacency failures originate in the '
  'planning layer: the picker has no notion of '
  'neighbours and approaches one target without '
  'considering whether the gripper width clears the '
  'next one. It is also worth noting that the current '
  'system does not act on the stall value at all — '
  'every PICK_ATTEMPT proceeds through the full place '
  'sequence regardless of how empty the grasp was. '
  'Adding a stall-threshold check after CLOSE_GRIPPER '
  '(for example, abort and re-detect if stall < 0.05 '
  'for tomato or strawberry) would catch every A-type '
  'failure at the source and convert it into a retry, '
  'with no hardware changes.')

H2('7.7 Aggregate Results')
P('Figure 4 summarises the overall picture. The '
  'autonomous baseline (Trials 1 and 2 combined) reached '
  '26 out of 27 (96.3 %), the autonomous robustness run '
  'reached 7 out of 14 (50.0 %), the pre-trial reached '
  '13 out of 18 (72.2 %), and the teleop test reached 0 '
  'out of 12 (0 %). The outcome composition shows that '
  'the baseline failures are almost entirely vision-'
  'related (the single F09 empty grasp), the robustness '
  'failures span all major failure categories '
  '(adjacency, inverted-pose, detection rejection), the '
  'pre-trial failures are dominated by physical (slip '
  'and empty) failures on the tomato, and the teleop '
  'failures are uniformly caused by the control-layer '
  'drift problem.')
FIG('figD_mode_comparison.png',
    'Figure 4: Aggregate task completion rates across '
    'the four test modes, with outcome composition.',
    width_in=5.5)

H1('8. Discussion')
P('The brief asks for a comparison between simulation '
  'results and physical-system performance, and for a '
  'discrepancy analysis. Because we did not have a '
  'QUARC license (Section 3.1) and could not run a '
  'Simulink dynamic model on the QArm directly, the '
  'Simulink work in Chapter 4 was a structural '
  'validation rather than a numerical dynamic '
  'comparison. Each Simulink slice was tested against '
  'its Python counterpart on identical inputs, and the '
  'output values agreed at machine precision. The '
  'physical performance numbers reported in Chapter 7 '
  'are therefore the source for the discrepancy '
  'analysis between expected and actual behaviour, and '
  'they take three forms.')
P('The first discrepancy is detection misclassification. '
  'The HSV gates and shape conditions described in '
  'Section 5.1 were tuned on canonical (upright, well-'
  'separated) fruit images. In Trial 1 the system '
  'misclassified one fruit out of fourteen; in Trial 2 '
  'it misclassified zero out of thirteen; and in '
  'Trial 3 it misclassified three out of fourteen. The '
  'two extreme cases — the tomato/strawberry confusion '
  'when the tomato is touching a banana, and the total '
  'rejection of strawberries rotated past 90° — both '
  'come from the structure of the strawberry '
  'acceptance path. The classical pipeline is fast and '
  'inspectable but brittle to non-canonical inputs, '
  'which is the well-known trade-off against learned '
  'detectors.')
P('The second discrepancy is the false-positive grasp. '
  'The autonomous controller logs a "Placed" event '
  'whenever the lift-and-release motion completes, '
  'whether or not the gripper actually retained the '
  'fruit. Section 7.6 quantifies this: empty-grasp '
  'events have stall around 0.02 and are easy to '
  'detect, but the controller never reads the stall '
  'and therefore treats them as successful pickups. '
  'This explains the gap between the per-attempt '
  'success rate (33/46 = 71.7 %) and the per-fruit '
  'task completion rate (33/41 = 80.5 %) and is '
  'directly responsible for the F09 failure in Trial 1 '
  'being invisible to the controller.')
P('The third discrepancy is the teleop control-layer '
  'defect. The IK call with gamma = 0.0 inside '
  '_teleop redistributes all four joints on every '
  'Cartesian step, so the end-effector drifts down in '
  'Z when the operator presses X+ or Y+. The gripper '
  'command is also anchored to the readback rather '
  'than to the last commanded value, so any slip is '
  'reinforced rather than corrected. Both are visible '
  'only when the autonomous wrapper is removed, and '
  'they explain the full 0/12 teleop result. The '
  'discrepancy here is not between sim and real — it '
  'is between the expected behaviour of the GUI '
  'controls and what the underlying control stack '
  'actually does.')
P('Taken together, these three discrepancies map onto '
  'the three failure classes defined in Section 7.6: '
  'misclassifications correspond to the perception '
  'layer, false-positive grasps correspond to the '
  'perception-to-action interface and the gripper '
  'hardware, and teleop failures correspond to the '
  'control layer.')

H1('9. Conclusion and Future Work')

H2('9.1 Conclusion')
P('The project delivered a working autonomous pick-and-'
  'place system on the Quanser QArm 4-DOF platform, '
  'paired with an Intel RealSense D415 overhead camera. '
  'The system identifies bananas, tomatoes, and '
  'strawberries in a workspace of fourteen fruits and '
  'sorts them into three labelled baskets. We measured '
  'end-to-end performance over four test modes and a '
  'total of 89 pick attempts on physical fruit. The '
  'autonomous baseline reached 26/27 = 96.3 % task '
  'completion under standard upright placement, with a '
  'cycle time of about 29 seconds per pick attempt; '
  'the autonomous robustness run reached 7/14 = 50.0 % '
  'under deliberate adversarial placement; the pre-'
  'trial reached 13/18 = 72.2 % over single-fruit '
  'fixed-position tests; the teleoperation test '
  'reached 0/12 = 0 %. The 96 % autonomous baseline '
  'was achieved under the realistic operational '
  'envelope for a sorting robot of this design, and '
  'the system was reliable, repeatable across two '
  'trials, and fast enough that the cycle time would '
  'not be the limiting factor for any practical '
  'throughput target. Within this operational '
  'envelope, we therefore consider the original goal '
  '— autonomously picking and placing labelled fruit '
  'on the QArm — to be met.')
P('The robustness test, the pre-trial, and the teleop '
  'test together drew a clear map of where the system '
  'stops working. Three weak points were identified. '
  'First, the rigid two-finger gripper does not '
  'provide enough friction for smooth and heavy '
  'fruits, which made the tomato the worst-performing '
  'fruit class in every mode. Second, the grasp pose '
  'estimation is based on fixed type-specific biases '
  'tuned for upright fruits, and these biases land '
  'the gripper outside the actual contact surface '
  'when the fruit is inverted or sits next to another '
  'fruit. Third, the teleop layer has two control-'
  'architecture defects that together prevent any '
  'human operator from completing the task without a '
  'redesign.')

H2('9.2 Future Work')
P('We propose five concrete improvements that would '
  'directly target the weaknesses observed in the '
  'test data, in approximate order of expected impact '
  'per hour of work.')
P('Soft or compliant gripper. The single highest-'
  'value mechanical change is replacing the rigid '
  'two-finger gripper with a compliant design — for '
  'example a 3D-printed silicone-padded fingertip '
  'insert or a fully soft pneumatic gripper. Every '
  'B-type slip-during-ascent failure we observed (the '
  'P1 pre-trial tomato is the canonical example) '
  'would be reduced by even a modest increase in '
  'finger-to-fruit friction, and the same change '
  'would also reduce the geometric sensitivity of the '
  'grasp because a compliant pad conforms to the '
  'fruit shape rather than requiring the visual '
  'centroid to be precisely on the contact axis.')
P('Stall-based closed-loop pick verification. The '
  'system currently treats every PICK_ATTEMPT as '
  'successful and proceeds through the full lift-and-'
  'place sequence regardless of how empty the grasp '
  'was. Adding a single stall-threshold check after '
  'the CLOSE_GRIPPER step (for example, abort and '
  're-detect if stall < 0.05 for tomato or '
  'strawberry) would catch every A-type empty-grasp '
  'failure at the source and convert it into a retry '
  'rather than a wasted cycle. This is a small code '
  'change with no hardware cost.')
P('Point-cloud-based grasp pose estimation. The '
  'current pipeline computes a 2D centroid from the '
  'colour image and then applies fixed type-specific '
  'biases to produce a 3D grasp point. This breaks '
  'under non-canonical orientations because the '
  'biases were tuned for upright fruit only. '
  'Replacing the fixed-bias step with a depth-driven '
  'grasp pose estimator (for example, fitting an '
  'oriented bounding box to the fruit\'s point cloud '
  'and selecting the gripper approach axis along the '
  'shortest principal component) would generalise '
  'across orientations without per-class tuning. The '
  'D415 already provides the depth data; only the '
  'picker logic needs to change.')
P('Learned detector for non-canonical orientations. '
  'The current OpenCV detector relies on hand-tuned '
  'HSV ranges, contour shape, and a green-calyx '
  'signal to classify fruit, and we found two '
  'failure modes that are intrinsic to this approach: '
  'the strawberry rejection at greater than 90° '
  'rotation and the strawberry/tomato confusion when '
  'fruits are adjacent. A small CNN classifier '
  'trained on labelled images of upright, inverted, '
  'and rotated fruit would not have either failure '
  'mode. Even a low-data approach (fine-tuning a '
  'pretrained backbone on a few hundred lab images) '
  'would likely close both gaps.')
P('Teleop control rework. The 0/12 result is not a '
  'tuning issue; it requires a redesign of _teleop '
  'in main_gui.py. The two specific changes needed '
  'are: (a) replace joint-space interpolation under '
  'a fixed orientation constraint with a Cartesian-'
  'space straight-line trajectory, and (b) decouple '
  'the gripper setpoint from the readback so a slip '
  'does not propagate into the next teleop press. '
  'Both changes are localised to a single file and '
  'would convert the teleop layer from a control '
  'group into a usable mode.')

# --- 10. Feedback Reflection ---
PAGEBREAK()
H1('10. Reflection on Lab Feedback')
P('This one-page reflection summarises the key '
  'feedback we received from the lecturers and the '
  'PGTAs during the lab sessions, and how that '
  'feedback shaped the final design.')

P('Feedback 1 — Calibration approach. In an early '
  'lab session our hand-eye calibration was based on '
  'driving the arm to several known points and '
  'tracking a coloured ball with the D415. A PGTA '
  'pointed out that this approach gave us a '
  'nonlinear fit with very few constraints, and that '
  'any small movement of the arm base between '
  'calibration and use would invalidate the result. '
  'The point was that we were calibrating the wrong '
  'thing — what we actually needed was a flat '
  'reference that ties the camera image directly to '
  'the table plane, not a 3D point cloud through an '
  'unconstrained pose. Acting on this feedback, we '
  'rewrote the calibration pipeline around a 7×5 '
  'inner-corner chessboard with a touch-probed '
  'origin, which is the procedure described in '
  'Sections 5.1.3 and 7.1. This change reduced our '
  'setup time per session and made the system '
  'robust to the small bench movements that used to '
  'break us.')

P('Feedback 2 — Detection robustness. During an '
  'early demonstration run, our detector '
  'misclassified several strawberries as tomatoes '
  'whenever the calyx was partly occluded. The '
  'feedback we received was to think about each '
  'class in terms of "what evidence is positive for '
  'this class" rather than just confidence numbers, '
  'and to make sure each class had at least two '
  'independent ways to be accepted. This pushed us '
  'to replace a single confidence threshold with '
  'two independent positive signals on the '
  'strawberry path (a calyx green ratio above 0.05 '
  'OR a contour circularity below 0.4), and to '
  'extend the green-calyx search region beyond the '
  'bounding box. The Trial 3 detection-rejection '
  'failures (F13 and F14) are still a consequence '
  'of asking for both signals to fail '
  'simultaneously, but the same advice explains why '
  'the detector still recognised the inverted '
  'strawberries (F09, F11) — the bbox-padded green '
  'search caught the calyx even when the fruit was '
  'upside down, even though the grasp pose then '
  'failed for an unrelated reason.')

P('Feedback 3 — Test methodology. Our original '
  'test plan was a single autonomous run with all '
  'fourteen fruits in nominal placement. Feedback '
  'during a mid-project review was that a single '
  'mixed run would not let us attribute failures '
  'to specific layers — a tomato slip and a tomato '
  'misclassification would look the same in the '
  'aggregate score. The recommendation was to use '
  'a layered test design: a single-fruit pre-trial '
  'to isolate basic capability, a repeatability '
  'run to characterise the baseline, an '
  'adversarial run to probe boundaries, and a '
  'teleop control group. The four-mode structure '
  'used throughout Chapter 7 came directly from '
  'this advice, and as we noted in Section 7.7, it '
  'was the main reason we could separate '
  'perception failures from actuation failures in '
  'the final analysis.')

P('Reflection on the feedback process itself. '
  'Looking back, the most useful feedback was '
  'rarely a direct answer to a specific question; '
  'it was more often a reframing of the problem. '
  'The "calibrate the right thing" comment '
  'changed our entire vision pipeline; the "what '
  'evidence is positive" comment reshaped the '
  'detector logic; the "test design needs layers" '
  'comment defined the structure of Chapter 7. We '
  'did not always agree with the feedback at '
  'first — moving from a learned-looking 3D '
  'approach to a flat chessboard felt like a step '
  'backward at the time — but in each case the '
  'suggestion turned out to capture a constraint '
  'we had not noticed. For future projects we '
  'would schedule earlier and more frequent '
  'design reviews with the PGTAs, since most of '
  'the time saved by feedback came at the points '
  'where we changed direction, not at the points '
  'where we tuned parameters.')

# --- save ---
doc.save(OUT_PATH)
print('saved:', OUT_PATH)
